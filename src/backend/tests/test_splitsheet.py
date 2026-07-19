"""Tests for split sheet generation endpoint.

Covers:
- POST /splitsheet/generate  returns PDF StreamingResponse
- POST /splitsheet/generate  returns DOCX StreamingResponse
- POST /splitsheet/generate  returns 400 when no contributors
- POST /splitsheet/generate  returns 422 on invalid payload
"""

CONTRIBUTOR = {
    "name": "Jane Doe",
    "role": "Songwriter",
    "writer_share": 25.0,
    "publisher_share": 25.0,
    "ipi_number": "00123456789",
    "is_published": True,
    "publisher_name": "Indie Publishing",
    "publisher_ipi": "00987654321",
    "master_percentage": 50.0,
    "label": "Indie Label",
}

BASE_PAYLOAD = {
    "work_title": "Test Track",
    "work_type": "single",
    "split_type": "both",
    "date": "2026-04-10",
    "format": "pdf",
    "contributors": [CONTRIBUTOR],
}


class TestGenerateSplitSheetPDF:
    """POST /splitsheet/generate with format=pdf produces a valid PDF."""

    def test_returns_200_with_pdf_content_type(self, client, mock_supabase):
        """Successful PDF generation returns 200 and correct content-type."""
        response = client.post("/splitsheet/generate", json=BASE_PAYLOAD)

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/pdf"

    def test_response_body_is_non_empty(self, client, mock_supabase):
        """PDF body must be non-empty bytes."""
        response = client.post("/splitsheet/generate", json=BASE_PAYLOAD)

        assert response.status_code == 200
        assert len(response.content) > 0

    def test_content_disposition_uses_safe_filename(self, client, mock_supabase):
        """Content-Disposition header contains the sanitised work title and .pdf ext."""
        response = client.post("/splitsheet/generate", json=BASE_PAYLOAD)

        assert response.status_code == 200
        cd = response.headers.get("content-disposition", "")
        assert "Split_Sheet_" in cd
        assert ".pdf" in cd

    def test_multiple_contributors(self, client, mock_supabase):
        """Works with multiple contributors."""
        contributor_2 = {
            "name": "John Smith",
            "role": "Producer",
            "is_published": False,
            "publishing_share": 50.0,
            "master_percentage": 50.0,
        }
        payload = {**BASE_PAYLOAD, "contributors": [CONTRIBUTOR, contributor_2]}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200
        assert len(response.content) > 0

    def test_special_characters_in_title_are_sanitised(self, client, mock_supabase):
        """Special characters in work_title are replaced with underscores in filename."""
        payload = {**BASE_PAYLOAD, "work_title": "My Song: Version #2 (Remix)"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200
        cd = response.headers.get("content-disposition", "")
        # Colons, hashes, parens and spaces should be sanitised
        assert ":" not in cd
        assert "#" not in cd

    def test_minimal_contributor_fields(self, client, mock_supabase):
        """Contributor with only required fields (name, role) succeeds."""
        minimal_contributor = {"name": "Solo Artist", "role": "Composer"}
        payload = {**BASE_PAYLOAD, "contributors": [minimal_contributor]}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200


class TestGenerateSplitSheetDOCX:
    """POST /splitsheet/generate with format=docx produces a valid DOCX."""

    def test_returns_200_with_docx_content_type(self, client, mock_supabase):
        """Successful DOCX generation returns 200 and correct content-type."""
        payload = {**BASE_PAYLOAD, "format": "docx"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200
        assert (
            response.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_response_body_is_non_empty(self, client, mock_supabase):
        """DOCX body must be non-empty bytes."""
        payload = {**BASE_PAYLOAD, "format": "docx"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200
        assert len(response.content) > 0

    def test_content_disposition_has_docx_extension(self, client, mock_supabase):
        """Content-Disposition header contains .docx extension."""
        payload = {**BASE_PAYLOAD, "format": "docx"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200
        cd = response.headers.get("content-disposition", "")
        assert ".docx" in cd


class TestDocxContent:
    """Directly exercise the DOCX generator to assert publishing/master rendering."""

    @staticmethod
    def _docx_text(buffer):
        from docx import Document

        doc = Document(buffer)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def test_master_only_omits_ipi(self):
        """A master-only sheet must never render IPI/CAE text."""
        from splitsheet.docx_generator import generate_split_sheet_docx

        contributors = [
            {
                "name": "Jane Doe",
                "role": "Producer",
                "ipi_number": "00123456789",  # should be ignored on master-only
                "master_percentage": 100.0,
                "label": "Sub Pop",
            }
        ]
        text = self._docx_text(
            generate_split_sheet_docx(
                work_title="Track",
                work_type="single",
                split_type="master",
                date="2026-07-19",
                contributors=contributors,
            )
        )
        assert "IPI" not in text
        assert "00123456789" not in text
        assert "Master Royalty %" in text
        assert "Sub Pop" in text  # optional label column renders

    def test_published_writer_renders_writer_and_publisher_shares(self):
        """A published writer shows both a writer's share and a distinct publisher recipient."""
        from splitsheet.docx_generator import generate_split_sheet_docx

        contributors = [
            {
                "name": "Jane Doe",
                "role": "Songwriter",
                "writer_share": 25.0,
                "publisher_share": 25.0,
                "ipi_number": "00123456789",
                "is_published": True,
                "publisher_name": "Big Publisher",
                "publisher_ipi": "00987654321",
            }
        ]
        text = self._docx_text(
            generate_split_sheet_docx(
                work_title="Track",
                work_type="single",
                split_type="publishing",
                date="2026-07-19",
                contributors=contributors,
            )
        )
        assert "Writer's Share" in text
        assert "Publisher's Share" in text
        assert "Big Publisher" in text
        assert "IPI/CAE #00123456789" in text  # writer IPI in parties prose
        assert "self-published" not in text.lower()

    def test_self_published_writer_labeled(self):
        """A self-published writer is labeled as such, with no separate publisher."""
        from splitsheet.docx_generator import generate_split_sheet_docx

        contributors = [
            {
                "name": "Solo Writer",
                "role": "Composer",
                "is_published": False,
                "publishing_share": 100.0,
            }
        ]
        text = self._docx_text(
            generate_split_sheet_docx(
                work_title="Track",
                work_type="single",
                split_type="publishing",
                date="2026-07-19",
                contributors=contributors,
            )
        )
        assert "self-published" in text.lower()
        assert "Self-published" in text  # publisher column value
        assert "100.00%" in text  # their full publishing share renders


class TestGenerateSplitSheetValidation:
    """Input validation for POST /splitsheet/generate."""

    def test_returns_400_when_contributors_empty(self, client, mock_supabase):
        """Returns 400 when contributors list is empty."""
        payload = {**BASE_PAYLOAD, "contributors": []}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 400
        assert "contributor" in response.json()["detail"].lower()

    def test_returns_422_when_contributors_missing(self, client, mock_supabase):
        """Returns 422 when contributors field is absent."""
        payload = {k: v for k, v in BASE_PAYLOAD.items() if k != "contributors"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 422

    def test_returns_422_when_work_title_missing(self, client, mock_supabase):
        """Returns 422 when work_title is absent."""
        payload = {k: v for k, v in BASE_PAYLOAD.items() if k != "work_title"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 422

    def test_returns_422_when_date_missing(self, client, mock_supabase):
        """Returns 422 when date is absent."""
        payload = {k: v for k, v in BASE_PAYLOAD.items() if k != "date"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 422

    def test_defaults_format_to_pdf_when_not_specified(self, client, mock_supabase):
        """Omitting format defaults to PDF generation."""
        payload = {k: v for k, v in BASE_PAYLOAD.items() if k != "format"}

        response = client.post("/splitsheet/generate", json=payload)

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/pdf"


# ---------------------------------------------------------------------------
# Subscription gate tests
# ---------------------------------------------------------------------------


class TestSplitSheetGated:
    """POST /splitsheet/generate with a Free user at cap returns 402."""

    def test_generate_at_cap_full_path(self, client, mock_supabase):
        """Full path: real EntitlementsService runs over wired supabase tables.

        Catches count-query regressions (wrong table/filter/etc.) that the monkeypatch
        test would miss. Free tier max_split_sheets_per_month=5; user has 5 → 402.
        """
        from unittest.mock import MagicMock

        from tests.conftest import TEST_USER_ID, MockQueryBuilder

        FREE_TIER = {
            "tier": "free",
            "max_artists": 3,
            "max_projects": 3,
            "max_boards": -1,
            "max_tasks": 50,
            "max_storage_bytes": 1073741824,
            "max_split_sheets_per_month": 5,
            "max_oneclick_runs_per_month": 1,
            "zoe_enabled": False,
            "oneclick_enabled": True,
            "registry_enabled": False,
            "integrations_allowed": ["google_drive"],
            "updated_at": "2026-05-09T00:00:00+00:00",
        }
        FREE_SUB = {
            "id": "s1",
            "user_id": TEST_USER_ID,
            "tier": "free",
            "status": "active",
            "stripe_customer_id": None,
            "stripe_subscription_id": None,
            "stripe_price_id": None,
            "current_period_start": None,
            "current_period_end": None,
            "cancel_at_period_end": False,
            "canceled_at": None,
            "created_at": "2026-05-01T00:00:00+00:00",
            "updated_at": "2026-05-01T00:00:00+00:00",
        }
        AT_CAP_USAGE = {
            "user_id": TEST_USER_ID,
            "total_storage_bytes": 0,
            "split_sheets_this_period": 5,
            "zoe_queries_this_period": 0,
            "oneclick_runs_this_period": 0,
            "period_start": "2026-05-09T00:00:00+00:00",
            "period_end": "2099-05-09T00:00:00+00:00",
            "updated_at": "2026-05-09T00:00:00+00:00",
        }

        # Reset the EntitlementsService singleton so it is rebuilt with free-tier data
        import subscriptions.deps as _sub_deps

        _sub_deps._entitlements_service = None

        def _table(name):
            b = MockQueryBuilder()
            if name == "subscriptions":
                b.execute.return_value = MagicMock(data=[FREE_SUB], count=1)
            elif name == "tier_entitlements":
                b.execute.return_value = MagicMock(data=[FREE_TIER], count=1)
            elif name == "tier_overrides":
                b.execute.return_value = MagicMock(data=[], count=0)
            elif name == "usage_counters":
                b.execute.return_value = MagicMock(data=[AT_CAP_USAGE], count=1)
            return b

        mock_supabase.table.side_effect = _table

        VALID_BODY = {
            "work_title": "Test Song",
            "date": "2026-05-10",
            "contributors": [
                {"name": "Alice", "role": "Songwriter"},
                {"name": "Bob", "role": "Songwriter"},
            ],
        }
        resp = client.post("/splitsheet/generate", json=VALID_BODY)
        assert resp.status_code == 402
        assert "split sheet" in resp.json()["detail"].lower()

    def test_generate_increments_usage_counter_on_success(self, client, mock_supabase, monkeypatch):
        """Successful split-sheet generation increments split_sheets_this_period.

        Verifies that increment_usage is called exactly once with the correct
        arguments on a successful generation, catching regressions where the
        increment is accidentally removed or called with wrong args.
        """
        from unittest.mock import MagicMock

        from tests.conftest import TEST_USER_ID

        # Pro user — mock_supabase already defaults to Pro so the gate passes.
        # Patch _get_entitlements_service at the splitsheet router module level
        # so the mock's increment_usage is the one that gets called.
        mock_svc = MagicMock()
        mock_svc.can.return_value = MagicMock(allowed=True)

        import splitsheet.router as ssrouter

        monkeypatch.setattr(ssrouter, "_get_entitlements_service", lambda: mock_svc)

        VALID_BODY = {
            "work_title": "Test Song",
            "date": "2026-05-10",
            "contributors": [
                {"name": "Alice", "role": "Songwriter"},
                {"name": "Bob", "role": "Songwriter"},
            ],
        }
        resp = client.post("/splitsheet/generate", json=VALID_BODY)

        # Generation succeeded
        assert resp.status_code == 200, f"Expected 200, got {resp.status_code}: {resp.text}"

        # Counter incremented exactly once with the correct args
        mock_svc.increment_usage.assert_called_once_with(
            TEST_USER_ID,
            "split_sheets_this_period",
        )

    def test_generate_at_cap_returns_402(self, client, mock_supabase, monkeypatch):
        """Free user at 5/5 split sheets this period → POST /splitsheet/generate returns 402."""
        from unittest.mock import MagicMock

        from subscriptions import enforcement
        from subscriptions.models import CheckResult

        # Patch enforcement._service to return a service that denies GENERATE_SPLIT_SHEET
        deny_result = CheckResult(
            allowed=False,
            reason="You've used your 5 split sheet(s) for this period.",
            upgrade_required=True,
        )
        svc = MagicMock()
        svc.can.return_value = deny_result
        monkeypatch.setattr(enforcement, "_service", lambda: svc)

        VALID_BODY = {
            "work_title": "Test Song",
            "date": "2026-05-10",
            "contributors": [
                {"name": "Alice", "role": "Songwriter"},
                {"name": "Bob", "role": "Songwriter"},
            ],
        }
        resp = client.post("/splitsheet/generate", json=VALID_BODY)

        assert resp.status_code == 402
        assert "split sheet" in resp.json()["detail"].lower()
