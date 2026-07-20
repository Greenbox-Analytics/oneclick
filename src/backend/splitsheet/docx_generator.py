import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def generate_split_sheet_docx(
    work_title: str,
    work_type: str,
    split_type: str,
    date: str,
    contributors: list[dict],
) -> io.BytesIO:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    brand_color = RGBColor(0x1A, 0x3A, 0x2A)
    gray_color = RGBColor(0x55, 0x55, 0x55)

    needs_pub = split_type in ("publishing", "both")
    needs_master = split_type in ("master", "both")

    work_type_display = work_type.upper() if work_type else "SINGLE"

    def _add_shading(cell, fill_color):
        shading_element = cell._element.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd", {}
        )
        shading_element.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", fill_color)
        shading_element.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "clear")
        tcPr = cell._element.get_or_add_tcPr()
        tcPr.append(shading_element)

    # --- Header ---
    title = doc.add_heading("SPLIT SHEET AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = brand_color
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph("Music Work Royalty Split Agreement")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = gray_color
        run.font.size = Pt(12)

    doc.add_paragraph("_" * 85)

    # --- Section 1: Parties ---
    heading = doc.add_heading("Section 1: Parties to This Agreement", level=2)
    for run in heading.runs:
        run.font.color.rgb = brand_color

    p = doc.add_paragraph()
    run = p.add_run(
        f'This Split Sheet Agreement (the "Agreement") is entered into as of {date}, '
        f"by and between the following parties:"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    for c in contributors:
        name = c.get("name", "")
        role = c.get("role", "")

        p = doc.add_paragraph()
        run = p.add_run(f"{name}")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f' (hereinafter referred to as "{role}")')
        run.font.size = Pt(10)
        # IPI and publisher affiliation only apply to the publishing (composition) side.
        if needs_pub:
            ipi = c.get("ipi_number", "") or ""
            if ipi:
                run = p.add_run(f", IPI/CAE #{ipi}")
                run.font.size = Pt(10)
            if c.get("is_published"):
                publisher_name = c.get("publisher_name", "") or "their publisher"
                suffix = f", published by {publisher_name}"
                publisher_ipi = c.get("publisher_ipi", "") or ""
                if publisher_ipi:
                    suffix += f" (IPI/CAE #{publisher_ipi})"
                run = p.add_run(suffix)
                run.font.size = Pt(10)
            else:
                run = p.add_run(", self-published")
                run.font.size = Pt(10)
        elif c.get("label"):
            run = p.add_run(f", {c.get('label')}")
            run.font.size = Pt(10)

    doc.add_paragraph()

    # --- Section 2: Musical Work ---
    heading = doc.add_heading("Section 2: Musical Work", level=2)
    for run in heading.runs:
        run.font.color.rgb = brand_color

    p = doc.add_paragraph()
    run = p.add_run("The parties agree that this Agreement pertains to the following musical work:")
    run.font.size = Pt(10)

    doc.add_paragraph()

    details = [
        ("Work Title:", work_title),
        ("Work Type:", f"{work_type_display} (song)"),
        ("Date:", date),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}  ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(value)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # --- Section 3: Royalty Splits ---
    heading = doc.add_heading("Section 3: Royalty Percentage Splits", level=2)
    for run in heading.runs:
        run.font.color.rgb = brand_color

    def _render_table(headers, body_rows, total_row_values):
        """Render a styled table: brand header, zebra body rows, bold TOTAL row.

        `body_rows` and `total_row_values` are lists of already-stringified cells.
        """
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_row = table.rows[0]
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _add_shading(cell, "1A3A2A")

        for idx, values in enumerate(body_rows, 1):
            row = table.add_row()
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            if idx % 2 == 0:
                for cell in row.cells:
                    _add_shading(cell, "F5F5F5")

        total_row = table.add_row()
        for i, val in enumerate(total_row_values):
            total_row.cells[i].text = val
        for cell in total_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        doc.add_paragraph()

    def build_publishing_section():
        """Publishing splits — writer's share and publisher's share per writer."""
        p = doc.add_paragraph()
        run = p.add_run("Publishing Royalties:")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(
            " Publishing income is divided into a writer's share and a publisher's "
            "share. The parties agree to the following publishing royalty percentage "
            f'splits for the work titled "{work_title}":'
        )
        run.font.size = Pt(10)

        doc.add_paragraph()

        for c in contributors:
            name = c.get("name", "")
            p = doc.add_paragraph(style="List Bullet")
            if c.get("is_published"):
                writer_pct = c.get("writer_share", 0) or 0
                publisher_pct = c.get("publisher_share", 0) or 0
                recipient = c.get("publisher_name", "") or "their publisher"
                run = p.add_run(
                    f"{name} shall receive {writer_pct:.2f}% as Writer's Share; "
                    f"{recipient} shall receive {publisher_pct:.2f}% as Publisher's Share."
                )
            else:
                pub_pct = c.get("publishing_share", 0) or 0
                run = p.add_run(f"{name} shall receive {pub_pct:.2f}% of all publishing royalties (self-published).")
            run.font.size = Pt(10)

        doc.add_paragraph()

        body_rows = []
        writer_total = 0.0
        publisher_total = 0.0
        for c in contributors:
            if c.get("is_published"):
                writer_pct = c.get("writer_share", 0) or 0
                publisher_pct = c.get("publisher_share", 0) or 0
                publisher = c.get("publisher_name", "") or ""
            else:
                writer_pct = c.get("publishing_share", 0) or 0
                publisher_pct = 0
                publisher = "Self-published"
            writer_total += writer_pct
            publisher_total += publisher_pct
            body_rows.append(
                [
                    c.get("name", ""),
                    c.get("role", ""),
                    f"{writer_pct:.2f}%",
                    publisher,
                    f"{publisher_pct:.2f}%",
                ]
            )

        _render_table(
            ["Writer", "Role", "Writer's Share %", "Publisher", "Publisher's Share %"],
            body_rows,
            ["", "TOTAL", f"{writer_total:.2f}%", "", f"{publisher_total:.2f}%"],
        )

    def build_master_section():
        """Master splits — sound recording ownership. No IPI, no publisher share."""
        has_label = any((c.get("label") or "").strip() for c in contributors)

        p = doc.add_paragraph()
        run = p.add_run("Master Royalties:")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(
            f' The parties agree to the following master royalty percentage splits for the work titled "{work_title}":'
        )
        run.font.size = Pt(10)

        doc.add_paragraph()

        for c in contributors:
            name = c.get("name", "")
            role = c.get("role", "")
            pct = c.get("master_percentage", 0) or 0
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f'{name} ("{role}") shall receive {pct:.2f}% of all master royalties.')
            run.font.size = Pt(10)

        doc.add_paragraph()

        body_rows = []
        for c in contributors:
            pct = c.get("master_percentage", 0) or 0
            if has_label:
                body_rows.append([c.get("name", ""), c.get("role", ""), c.get("label", "") or "", f"{pct:.2f}%"])
            else:
                body_rows.append([c.get("name", ""), c.get("role", ""), f"{pct:.2f}%"])

        total_pct = sum((c.get("master_percentage", 0) or 0) for c in contributors)
        if has_label:
            headers = ["Party Name", "Role", "Label / Master Owner", "Master Royalty %"]
            total_row_values = ["", "TOTAL", "", f"{total_pct:.2f}%"]
        else:
            headers = ["Party Name", "Role", "Master Royalty %"]
            total_row_values = ["", "TOTAL", f"{total_pct:.2f}%"]

        _render_table(headers, body_rows, total_row_values)

    if needs_pub:
        build_publishing_section()

    if needs_master:
        build_master_section()

    # --- Section 4: Signatures ---
    heading = doc.add_heading("Section 4: Signatures", level=2)
    for run in heading.runs:
        run.font.color.rgb = brand_color

    p = doc.add_paragraph()
    run = p.add_run(
        "By signing below, each party acknowledges and agrees to the royalty "
        "percentage splits as outlined in Section 3 of this Agreement."
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    for c in contributors:
        name = c.get("name", "")
        role = c.get("role", "")

        p = doc.add_paragraph()
        run = p.add_run(f"Name: {name}")
        run.font.size = Pt(10)
        p.add_run("          ")
        run = p.add_run(f"Role: {role}")
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run("Signature: ___________________________")
        run.font.size = Pt(10)
        p.add_run("          ")
        run = p.add_run("Date: _______________")
        run.font.size = Pt(10)

        doc.add_paragraph()

    # --- Disclaimer ---
    doc.add_paragraph("_" * 85)
    disclaimer = doc.add_paragraph(
        "This split sheet agreement documents the agreed-upon royalty ownership percentages "
        "for the above-referenced musical work. All parties acknowledge and agree to the "
        "royalty splits as outlined. This document is not a substitute for a legally binding "
        "contract and all parties are advised to seek independent legal counsel."
    )
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.italic = True

    footer = doc.add_paragraph(f"Generated by Msanii \u2014 {date}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
